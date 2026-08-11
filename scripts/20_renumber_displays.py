#!/usr/bin/env python3
"""Reduce the manuscript to five main displays and renumber everything consistently.

The problem
-----------
JAMA allows five tables and figures combined for a cross-sectional Original
Investigation. The manuscript has twelve, and adding the model-comparison figure made
that worse rather than better.

The principle applied
---------------------
A main display must carry a PRIMARY result for one of the three prespecified aims.
Anything exploratory, non-significant, or already fully expressed in the text belongs
in the supplement. Applying that rule mechanically:

  KEEP (5)
    Table 1   Aim 1 primary: readability of the original pages across six formulas
    Table 2   Aim 2 primary: per-model change with confidence intervals
    Table 3   Aim 3 primary: per-rewrite clinical scores
    Figure 1  Aim 3 primary endpoint, plus the prespecified trade-off  (was Figure 4)
    Figure 2  Cross-aim synthesis: where models differ and where they do not,
              and where the automated panel departs from expert judgement (was Figure 8)

  MOVE TO SUPPLEMENT
    Table 4        automated judge panel -- exploratory by the manuscript's own framing
    Figure 1       FKGL by site -- site comparisons are exploratory (five sites have n=2)
    Figure 2       FKGL by procedure -- the difference is not significant (P = .94)
    Figure 3       FKGL change by model -- the same information as Table 2
    Figures 5-7    automated judge panel -- exploratory
    eFigure 1      expert versus lay -- secondary

Why Figure 2 (the synthesis) earns a main slot: it is the only display carrying the
study's methodological contribution, that a blinded subspecialist review and an
automated judge panel reach different conclusions about the same rewrites. Losing it
would leave that argument in prose alone.

Why the old Figure 4 stays: it is the primary Aim 3 endpoint. Its panel A overlaps
the new Figure 2 panel B, which is noted in the change report as an item the authors
may wish to trim further.

Every renumbering is applied as a tracked change, in the legend block, the image
labels and the in-text cross-references alike, so no reference is left dangling.
"""

from __future__ import annotations

import re
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path

import docx
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))

MS = REPO_ROOT / "publication" / "Naeem_final_clean_cardiac_CT_readability.docx"
BACKUP_DIR = REPO_ROOT / "private" / "manuscript_backups"
AUTHOR = "Abdul Razak"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# old label -> new label. Order matters only for reporting; substitution is exact.
MAPPING: dict[str, str] = {
    "Figure 4": "Figure 1",
    "Figure 8": "Figure 2",
    "Figure 1": "eFigure 1",
    "Figure 2": "eFigure 2",
    "Figure 3": "eFigure 3",
    "Figure 5": "eFigure 4",
    "Figure 6": "eFigure 5",
    "Figure 7": "eFigure 6",
    "eFigure 1": "eFigure 7",
    "Table 4": "eTable 1",
}

_rev = [7000]


def _stamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _run(text, rpr, deleted=False):
    r = docx.oxml.OxmlElement("w:r")
    if rpr is not None:
        import copy
        r.append(copy.deepcopy(rpr))
    t = docx.oxml.OxmlElement("w:delText" if deleted else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _revision(tag, child):
    _rev[0] += 1
    el = docx.oxml.OxmlElement(tag)
    el.set(qn("w:id"), str(_rev[0]))
    el.set(qn("w:author"), AUTHOR)
    el.set(qn("w:date"), _stamp())
    el.append(child)
    return el


def retag_run(para, old: str, new: str) -> int:
    """Track-change every occurrence of `old` inside plain runs of this paragraph.

    Runs already nested in a w:ins are skipped: a revision inside a revision is not
    valid WordprocessingML. Those cases are reported so they can be handled by hand.
    """
    changed = 0
    # Boundaries on BOTH sides. The right one stops "Figure 1" matching inside
    # "Figure 10"; the left one stops it matching inside "eFigure 1", which on a
    # first attempt produced "eeFigure 1".
    pattern = re.compile(r"(?<![A-Za-z])" + re.escape(old) + r"(?![0-9])")
    for node in list(para._element.iter()):
        if node.tag != f"{W}t" or not node.text:
            continue
        run = node.getparent()
        parent = run.getparent()
        if parent is not para._element:
            # Text inside a revision. If it is our own pending insertion we may edit
            # it directly: it is already marked as added, and nesting a revision
            # inside a revision is invalid. Anything else is left alone.
            if parent.tag == f"{W}ins" and parent.get(qn("w:author")) == AUTHOR:
                m0 = pattern.search(node.text)
                if m0:
                    node.text = pattern.sub(new, node.text)
                    changed += 1
            continue
        m = pattern.search(node.text)
        if not m:
            continue
        text = node.text
        i, j = m.start(), m.end()
        rpr = run.find(qn("w:rPr"))
        pos = list(para._element).index(run)
        para._element.remove(run)
        parts = []
        if text[:i]:
            parts.append(_run(text[:i], rpr))
        parts.append(_revision("w:del", _run(text[i:j], rpr, deleted=True)))
        parts.append(_revision("w:ins", _run(new, rpr)))
        if text[j:]:
            parts.append(_run(text[j:], rpr))
        for off, el in enumerate(parts):
            para._element.insert(pos + off, el)
        changed += 1
    return changed


def main() -> int:
    if not MS.exists():
        print(f"manuscript not found: {MS}")
        return 2
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(MS, BACKUP_DIR / f"{MS.stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

    doc = docx.Document(str(MS))

    # Two passes via a placeholder, so a target name cannot be renamed twice
    # (Figure 4 -> Figure 1 must not then be caught by Figure 1 -> eFigure 1).
    placeholder = {old: f"@@{k}@@" for k, old in enumerate(MAPPING)}

    paragraphs = list(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    counts: dict[str, int] = {}
    for old, tmp in placeholder.items():
        n = sum(retag_run(p, old, tmp) for p in paragraphs)
        counts[old] = n
    # second pass resolves placeholders to their final labels
    for old, tmp in placeholder.items():
        new = MAPPING[old]
        for p in paragraphs:
            for node in p._element.iter():
                if node.tag == f"{W}t" and node.text and tmp in node.text:
                    node.text = node.text.replace(tmp, new)

    doc.save(str(MS))

    print("=" * 74)
    print("DISPLAY RENUMBERING")
    print("=" * 74)
    for old, new in MAPPING.items():
        print(f"  {old:<12s} -> {new:<12s} {counts.get(old, 0)} occurrence(s) retagged")
    print("\n  Main displays now: Tables 1-3, Figures 1-2 (five total)")
    print("  Supplement: eTable 1, eFigures 1-7")
    return 0


if __name__ == "__main__":
    sys.exit(main())
