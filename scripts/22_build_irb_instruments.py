#!/usr/bin/env python3
"""Build the IRB submission copies of the review instrument, in Word and PDF.

An IRB needs to see exactly what a participant saw. The instrument reviewers used was
an HTML page, which is not a submittable artifact, so this renders the same content as
a paginated document in both formats an IRB office will accept.

Two variants are produced, because reviewers scored two:

  labeled   the standard instrument, which names the two passages "Original page" and
            "AI rewrite"
  neutral   identical items and identical order, with wording that never frames the
            task as original-versus-AI ("Reference passage" / "Passage to score")

The distinction matters to an IRB: the neutral variant exists precisely so that some
reviewers could score without knowing which passage was machine-generated, and the
submission should show both.

Content comes from the same committed packet the reviewers scored
(`blinded_review_packet_with_text.csv`), so the documents cannot drift from what was
actually administered. Nothing is retyped.

Participant confidentiality: the instrument carries a participant ID field, never a
name, matching how the score sheets are stored.
"""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))

from src.config import REVIEW_DIR  # noqa: E402

OUT = REPO_ROOT / "IRB" / "instruments"
PACKET = REVIEW_DIR / "blinded_review_packet_with_text.csv"

STUDY = "Readability of Online Patient-Education Materials for Pre-Procedure Cardiac CT"
SUBTITLE = "Clinical Accuracy Review Instrument"

MANUSCRIPT = REPO_ROOT / "publication" / "Naeem_final_clean_cardiac_CT_readability.docx"

# Fallback used only when the manuscript is unavailable (it is gitignored, so a fresh
# clone will not have it). Keeping the manuscript as the primary source means the
# investigator block cannot drift from the paper.
_FALLBACK_PI = [
    "Muhammad Naeem, MBBS, MD",
    "Division of Cardiothoracic Imaging, Department of Radiology",
    "Mayo Clinic AZ",
    "5777 East Mayo Blvd,",
    "Phoenix, AZ, 85054",
    "Email: naeem.muhammad@mayo.edu",
]


def principal_investigator() -> list[str]:
    """The investigator block, read from the manuscript's corresponding-author section.

    Reading it rather than retyping it means the address on an IRB submission and the
    address in the paper cannot disagree.
    """
    if not MANUSCRIPT.exists():
        return _FALLBACK_PI
    try:
        import docx as _docx

        from src.manuscript import paragraph_texts
    except ImportError:
        return _FALLBACK_PI
    paras = [p.strip() for p in paragraph_texts(_docx.Document(str(MANUSCRIPT))) if p.strip()]
    try:
        i = paras.index("Corresponding Author")
    except ValueError:
        return _FALLBACK_PI
    block = []
    for line in paras[i + 1:]:
        block.append(line)
        if line.lower().startswith("email:"):
            break
        if len(block) >= 8:
            break
    return block or _FALLBACK_PI

INK = RGBColor(0x1F, 0x2D, 0x3A)
ACCENT = RGBColor(0x2F, 0x5D, 0x8A)
PDF_INK = colors.HexColor("#1f2d3a")
PDF_ACCENT = colors.HexColor("#2f5d8a")
PDF_RULE = colors.HexColor("#d6dfe7")
PDF_TINT = colors.HexColor("#f4f7fa")

VARIANTS = {
    "labeled": {
        "name": "Standard (Labeled) Instrument",
        "left": "Original page",
        "right": "AI rewrite",
        "note": ("This variant identifies which passage is the original page and which is the "
                 "AI-generated rewrite. Reviewers remained blinded to which of the three "
                 "language models produced each rewrite."),
    },
    "neutral": {
        "name": "Neutral-Presentation Instrument",
        "left": "Reference passage",
        "right": "Passage to score",
        "note": ("This variant is identical in items, order and scoring scales, but its wording "
                 "never indicates which passage is machine-generated. It was administered so "
                 "that some reviewers scored without that knowledge."),
    },
}

SCALES = [
    ("Factual accuracy", "1–5",
     "Is every clinical statement in the passage correct? 5 = no factual errors; "
     "1 = serious factual error that could mislead a patient."),
    ("Completeness", "1–5",
     "Does the passage retain the clinically important content of the reference passage? "
     "5 = nothing important lost; 1 = major clinically important omission."),
    ("Added errors", "1–5",
     "Does the passage introduce assertions not supported by the reference passage or "
     "inconsistent with standard of care? 1 = none added; 5 = substantial invented content. "
     "Lower is better on this scale."),
]

INSTRUCTIONS = [
    "You will see a series of paired passages about a cardiac CT procedure. For each pair, "
    "read both passages and score the second one on the three scales below.",
    "Score each passage on its own merits. Do not compare passages across items, and do not "
    "revise an earlier score after seeing a later item.",
    "If a passage omits something you would expect but that is also absent from the reference "
    "passage, this is not an omission by the passage under review.",
    "Record your participant ID at the top of the response sheet. Do not write your name on "
    "any page of this instrument.",
]


# --------------------------------------------------------------------------- Word

def _shade(cell, hex_fill: str) -> None:
    el = cell._tc.get_or_add_tcPr()
    shd = el.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    el.append(shd)


def build_docx(variant: str, items: pd.DataFrame, path: Path) -> None:
    cfg = VARIANTS[variant]
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.9)
        s.left_margin = s.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    def para(text, size=10.5, bold=False, color=INK, align=None, space_after=6, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = color
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        return p

    # --- cover ---
    para(STUDY, size=17, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(SUBTITLE, size=13, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(cfg["name"], size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    pi = principal_investigator()
    para("Principal Investigator", size=11, bold=True, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    for line in pi:
        para(line, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    doc.add_paragraph()

    meta = doc.add_table(rows=0, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in [("Document", "Review instrument as administered"),
                 ("Variant", cfg["name"]),
                 ("Items", f"{len(items)} paired passages"),
                 ("Scales", "3 ordinal scales, 1–5"),
                 ("Prepared", date.today().strftime("%d %B %Y")),
                 ("Participant identifier", "ID only; no names recorded")]:
        row = meta.add_row().cells
        row[0].text, row[1].text = k, v
        row[0].paragraphs[0].runs[0].bold = True
        for c in row:
            c.paragraphs[0].runs[0].font.size = Pt(10)
    meta.columns[0].width = Inches(2.1)
    meta.columns[1].width = Inches(4.0)

    doc.add_paragraph()
    para(cfg["note"], size=10, italic=True, color=INK, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.add_section(WD_SECTION.NEW_PAGE)

    # --- instructions ---
    para("Instructions to reviewers", size=13, bold=True, color=ACCENT, space_after=8)
    for line in INSTRUCTIONS:
        p = doc.add_paragraph(line, style="List Bullet")
        p.paragraph_format.space_after = Pt(5)

    doc.add_paragraph()
    para("Scoring scales", size=13, bold=True, color=ACCENT, space_after=8)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(("Scale", "Range", "Definition")):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        _shade(hdr[i], "EAF0F6")
    for name, rng, desc in SCALES:
        c = t.add_row().cells
        c[0].text, c[1].text, c[2].text = name, rng, desc
        for cell in c:
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    t.columns[0].width = Inches(1.35)
    t.columns[1].width = Inches(0.7)
    t.columns[2].width = Inches(4.6)

    doc.add_section(WD_SECTION.NEW_PAGE)

    # --- items ---
    para("Items", size=13, bold=True, color=ACCENT, space_after=10)
    for n, (_, row) in enumerate(items.iterrows(), start=1):
        para(f"Item {n} of {len(items)}    ·    Item ID {row.blind_id}",
             size=11, bold=True, color=ACCENT, space_after=6)

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        head = tbl.rows[0].cells
        for i, h in enumerate((cfg["left"], cfg["right"])):
            head[i].text = h
            head[i].paragraphs[0].runs[0].bold = True
            head[i].paragraphs[0].runs[0].font.size = Pt(10)
            _shade(head[i], "EAF0F6")
        body = tbl.add_row().cells
        for i, txt in enumerate((row.original_text, row.rewrite_text)):
            body[i].text = str(txt)
            for p in body[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
        tbl.columns[0].width = Inches(3.3)
        tbl.columns[1].width = Inches(3.3)

        doc.add_paragraph()
        resp = doc.add_table(rows=1, cols=4)
        resp.style = "Table Grid"
        rh = resp.rows[0].cells
        for i, h in enumerate(("Factual accuracy (1–5)", "Completeness (1–5)",
                               "Added errors (1–5)", "Comments")):
            rh[i].text = h
            rh[i].paragraphs[0].runs[0].bold = True
            rh[i].paragraphs[0].runs[0].font.size = Pt(9)
            _shade(rh[i], "F4F7FA")
        blank = resp.add_row().cells
        for c in blank:
            c.text = " "
            c.paragraphs[0].paragraph_format.space_after = Pt(14)

        if n < len(items):
            doc.add_page_break()

    doc.save(str(path))


# ---------------------------------------------------------------------------- PDF

def build_pdf(variant: str, items: pd.DataFrame, path: Path) -> None:
    cfg = VARIANTS[variant]
    ss = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=ss["Title"], fontSize=17, leading=21,
                        textColor=PDF_ACCENT, alignment=TA_CENTER, spaceAfter=4)
    h2 = ParagraphStyle("h2", parent=ss["Heading2"], fontSize=13, leading=16,
                        textColor=PDF_ACCENT, spaceBefore=4, spaceAfter=8)
    sub = ParagraphStyle("sub", parent=ss["Normal"], fontSize=12, leading=15,
                         alignment=TA_CENTER, textColor=PDF_INK, spaceAfter=2)
    body = ParagraphStyle("body", parent=ss["Normal"], fontSize=10, leading=13.5,
                          textColor=PDF_INK, alignment=TA_JUSTIFY, spaceAfter=5)
    small = ParagraphStyle("small", parent=body, fontSize=8.5, leading=11, alignment=0)
    item_h = ParagraphStyle("item", parent=ss["Heading3"], fontSize=11, leading=14,
                            textColor=PDF_ACCENT, spaceBefore=2, spaceAfter=6)

    def footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor("#7c8a97"))
        canvas.drawString(0.85 * inch, 0.55 * inch, f"{SUBTITLE} — {cfg['name']}")
        canvas.drawRightString(LETTER[0] - 0.85 * inch, 0.55 * inch, f"Page {doc_.page}")
        canvas.setStrokeColor(PDF_RULE)
        canvas.setLineWidth(0.5)
        canvas.line(0.85 * inch, 0.72 * inch, LETTER[0] - 0.85 * inch, 0.72 * inch)
        canvas.restoreState()

    doc = SimpleDocTemplate(str(path), pagesize=LETTER,
                            leftMargin=0.85 * inch, rightMargin=0.85 * inch,
                            topMargin=0.9 * inch, bottomMargin=0.9 * inch,
                            title=f"{SUBTITLE} — {cfg['name']}", author="Study team")
    pi_style = ParagraphStyle("pi", parent=ss["Normal"], fontSize=10.5, leading=13.5,
                              alignment=TA_CENTER, textColor=PDF_INK, spaceAfter=1)
    story = [Spacer(1, 0.85 * inch), Paragraph(STUDY, h1), Paragraph(SUBTITLE, sub),
             Paragraph(f"<b>{cfg['name']}</b>", sub), Spacer(1, 0.28 * inch),
             Paragraph("<b>Principal Investigator</b>", ParagraphStyle(
                 "pih", parent=pi_style, textColor=PDF_ACCENT, fontSize=11, spaceAfter=4))]
    story += [Paragraph(line, pi_style) for line in principal_investigator()]
    story += [Spacer(1, 0.28 * inch)]

    meta = [["Document", "Review instrument as administered"],
            ["Variant", cfg["name"]],
            ["Items", f"{len(items)} paired passages"],
            ["Scales", "3 ordinal scales, 1–5"],
            ["Prepared", date.today().strftime("%d %B %Y")],
            ["Participant identifier", "ID only; no names recorded"]]
    mt = Table([[Paragraph(f"<b>{a}</b>", small), Paragraph(b, small)] for a, b in meta],
               colWidths=[2.0 * inch, 4.0 * inch])
    mt.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LINEBELOW", (0, 0), (-1, -2), 0.4, PDF_RULE),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5), ("TOPPADDING", (0, 0), (-1, -1), 5)]))
    story += [mt, Spacer(1, 0.35 * inch), Paragraph(f"<i>{cfg['note']}</i>", body), PageBreak()]

    story += [Paragraph("Instructions to reviewers", h2)]
    for line in INSTRUCTIONS:
        story.append(Paragraph(f"•&nbsp;&nbsp;{line}", body))
    story += [Spacer(1, 0.18 * inch), Paragraph("Scoring scales", h2)]

    rows = [[Paragraph("<b>Scale</b>", small), Paragraph("<b>Range</b>", small),
             Paragraph("<b>Definition</b>", small)]]
    rows += [[Paragraph(n, small), Paragraph(r, small), Paragraph(d, small)]
             for n, r, d in SCALES]
    st = Table(rows, colWidths=[1.3 * inch, 0.65 * inch, 4.05 * inch], repeatRows=1)
    st.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eaf0f6")),
        ("GRID", (0, 0), (-1, -1), 0.4, PDF_RULE),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
    story += [st, PageBreak(), Paragraph("Items", h2)]

    for n, (_, row) in enumerate(items.iterrows(), start=1):
        # Passages are stacked rather than placed side by side. A reportlab table row
        # cannot split across pages, and several source pages run to thousands of
        # words, so a two-column layout overflows the frame. Stacked panels flow
        # naturally and read better for long clinical text.
        def panel(title, text):
            head = Table([[Paragraph(f"<b>{title}</b>", small)]], colWidths=[6.3 * inch])
            head.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#eaf0f6")),
                ("BOX", (0, 0), (-1, -1), 0.4, PDF_RULE),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
            out = [head, Spacer(1, 0.05 * inch)]
            for chunk in str(text).split("\n"):
                if chunk.strip():
                    safe = chunk.strip().replace("&", "&amp;").replace("<", "&lt;")
                    out.append(Paragraph(safe, small))
            out.append(Spacer(1, 0.12 * inch))
            return out

        story.append(Paragraph(
            f"Item {n} of {len(items)} &nbsp;·&nbsp; Item ID {row.blind_id}", item_h))
        story += panel(cfg["left"], row.original_text)
        story += panel(cfg["right"], row.rewrite_text)
        resp = Table(
            [[Paragraph("<b>Factual accuracy (1–5)</b>", small),
              Paragraph("<b>Completeness (1–5)</b>", small),
              Paragraph("<b>Added errors (1–5)</b>", small),
              Paragraph("<b>Comments</b>", small)],
             ["", "", "", ""]],
            colWidths=[1.45 * inch, 1.45 * inch, 1.45 * inch, 1.95 * inch], rowHeights=[None, 34])
        resp.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), PDF_TINT),
            ("GRID", (0, 0), (-1, -1), 0.4, PDF_RULE),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
        story += [Spacer(1, 0.06 * inch), resp]
        if n < len(items):
            story.append(PageBreak())

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def main() -> int:
    if not PACKET.exists():
        print(f"packet not found: {PACKET}")
        return 2
    items = pd.read_csv(PACKET)
    OUT.mkdir(parents=True, exist_ok=True)

    made = []
    for variant in VARIANTS:
        stem = f"Review_Instrument_{variant}"
        docx_path = OUT / f"{stem}.docx"
        pdf_path = OUT / f"{stem}.pdf"
        build_docx(variant, items, docx_path)
        build_pdf(variant, items, pdf_path)
        made += [docx_path, pdf_path]

    print("=" * 72)
    print("IRB REVIEW INSTRUMENTS")
    print("=" * 72)
    print(f"  items per document : {len(items)} paired passages")
    for p in made:
        print(f"  {p.relative_to(REPO_ROOT)}  ({p.stat().st_size / 1024:.0f} KB)")
    print("\n  Content is read from the committed review packet, so these cannot")
    print("  drift from the instrument the reviewers actually scored.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
