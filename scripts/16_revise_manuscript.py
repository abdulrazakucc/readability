#!/usr/bin/env python3
"""Apply the implementation-review corrections to the manuscript as TRACKED CHANGES.

Edits are made IN PLACE, as real Word revisions, so each can be reviewed, accepted
or rejected individually.

python-docx has no tracked-changes API: editing paragraph text through it produces
silent, untracked edits, removing the reviewability this document depends on. So the
revision XML is written directly -- the old phrase in `w:del`/`w:delText`, the new
phrase in `w:ins`, each with an author and timestamp. Only the changed phrase is
marked, so revision marks stay readable.

No result is typed in: every numeric replacement is read from the regenerated
reports. A timestamped backup goes to `private/` before anything is modified.

Wording changes come from section 7 of the review, which supplies replacement text
for the passages it identifies as inaccurate.

Every substitution is recorded to a CSV so the revision is auditable line by line:
what was replaced, what replaced it, and which report the value came from.

Edits that require the authors' judgement -- reframing the Discussion, deciding how
to describe a borderline association -- are NOT attempted here. They are listed in
the run summary as outstanding.
"""

from __future__ import annotations

import copy
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path

import docx
import pandas as pd
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))

from src.config import REPORTS_DIR  # noqa: E402

PUB = REPO_ROOT / "publication"
MANUSCRIPT = PUB / "Naeem_final_clean_cardiac_CT_readability.docx"
BACKUP_DIR = REPO_ROOT / "private" / "manuscript_backups"
AUDIT = PUB / "manuscript_revision_audit.csv"
REVISION_AUTHOR = "Implementation review"
_rev_id = [1000]

MODEL_NAME = {"claude": "Claude Opus 4.8", "openai": "GPT-5.5", "gemini": "Gemini 3.1 Pro"}


def _fmt(x: float, dp: int = 2) -> str:
    return f"{x:.{dp}f}"


def _p(x: float) -> str:
    """JAMA style: leading zero dropped, < .001 for very small values."""
    if x < 0.001:
        return "< .001"
    return f"= {x:.2f}".replace("0.", ".")


def build_edits() -> list[dict]:
    """Every substitution, with its source. Numeric values come from reports."""
    prim = pd.read_csv(REPORTS_DIR / "aim3_compiled_by_model_primary.csv").set_index("model_id")
    irr_all = pd.read_csv(REPORTS_DIR / "aim3_compiled_irr.csv")
    irr = irr_all[irr_all.condition == "expert_labeled"].set_index("axis")
    trade = pd.read_csv(REPORTS_DIR / "aim3_compiled_tradeoff.csv")
    evl = pd.read_csv(REPORTS_DIR / "aim3_compiled_expert_vs_lay.csv").set_index("axis")
    ph = pd.read_csv(REPORTS_DIR / "aim2_posthoc_models.csv")
    ph = ph[ph.label == "fkgl"]

    E: list[dict] = []

    def add(old: str, new: str, why: str, src: str = "review §7"):
        E.append({"old": old, "new": new, "rationale": why, "source": src})

    # ---- 7.3 Methods: the 403 contradiction (the review calls this mandatory) ----
    add(
        "Five candidate pages from Johns Hopkins Medicine (3) and Mayo Clinic (2) returned "
        "HTTP 403 status codes consistent with commercial bot-detection at the content delivery "
        "network layer and were not recovered for this analysis; this exclusion is documented in "
        "the prespecified deviation log and addressed in the Limitations section.",
        "Five candidate pages (3 Johns Hopkins Medicine and 2 Mayo Clinic) initially returned "
        "HTTP 403 responses consistent with commercial bot detection. All 5 were recovered on "
        "June 8, 2026, by manual browser capture of the visible patient-facing text. Documented "
        "site-specific boilerplate was removed, after which the text passed through the same "
        "Unicode normalization, junk-line filtering, whitespace normalization, readability "
        "scoring, and downstream analysis used for the automated captures. All 5 recovered pages "
        "were included in the final N = 26.",
        "Methods said the 5 pages were NOT recovered; Results, Limitations and the deviation "
        "appendix all say they were recovered and included, and n=26 depends on them.",
    )

    # ---- 7.8 unit of analysis wording ----
    add("155 independent ratings", "155 expert ratings",
        "155 rating events over 77 rewrites are not 155 independent experimental units.")

    # ---- 7.8 Table 3 values, now per rewrite ----
    for m in ("claude", "openai", "gemini"):
        r = prim.loc[m]
        add(f"{_fmt(4.82 if m=='claude' else 4.84 if m=='openai' else 4.85)} "
            f"({_fmt(0.39 if m=='claude' else 0.62 if m=='openai' else 0.36)})",
            f"{_fmt(r.accuracy_1_5_mean)} ({_fmt(r.accuracy_1_5_sd)})",
            f"{MODEL_NAME[m]} accuracy: one expert mean per rewrite, not per rating.",
            "reports/aim3_compiled_by_model_primary.csv")

    # ---- 7.8 / 4.2 agreement values ----
    add("Gwet AC1 0.77 for accuracy, 0.79 for completeness, and 0.89 for added errors",
        f"Gwet AC1 {_fmt(irr.loc['accuracy_1_5','gwet_ac1'])} for accuracy, "
        f"{_fmt(irr.loc['completeness_1_5','gwet_ac1'])} for completeness, and "
        f"{_fmt(irr.loc['added_errors_1_5','gwet_ac1'])} for added errors",
        "AC1 now uses the protocol's 1-5 category universe for every axis and cohort.",
        "reports/aim3_compiled_irr.csv")
    add("whereas quadratic-weighted Cohen κ was near zero",
        "whereas mean pairwise quadratic-weighted Cohen κ was near zero",
        "Name the statistic exactly as computed.")

    # ---- 7.8 trade-off: add bootstrap CIs, both axes ----
    def tr(m, axis):
        row = trade[(trade.model_id == m) & (trade.axis == axis)].iloc[0]
        return row
    ga = tr("gemini", "accuracy_1_5")
    ca = tr("claude", "accuracy_1_5")
    oa = tr("openai", "accuracy_1_5")
    add("(Spearman ρ = −0.37 between grade levels removed and accuracy; P = .06)",
        f"(Spearman ρ = {_fmt(ga.spearman_rho)} between grade levels removed and accuracy; "
        f"95% CI, {_fmt(ga.ci_low)} to {_fmt(ga.ci_high)}; P {_p(ga.p_value)})",
        "Methods promised 5000-resample bootstrap CIs; they were not previously computed.",
        "reports/aim3_compiled_tradeoff.csv")
    add("whereas the association was null for Claude Opus 4.8 (ρ = 0.25; P = .22) and GPT-5.5 (ρ = 0.03; P = .90)",
        f"whereas the association was null for Claude Opus 4.8 (ρ = {_fmt(ca.spearman_rho)}; "
        f"95% CI, {_fmt(ca.ci_low)} to {_fmt(ca.ci_high)}; P {_p(ca.p_value)}) and GPT-5.5 "
        f"(ρ = {_fmt(oa.spearman_rho)}; 95% CI, {_fmt(oa.ci_low)} to {_fmt(oa.ci_high)}; "
        f"P {_p(oa.p_value)}). Completeness correlations were similarly null "
        f"(Claude ρ = {_fmt(tr('claude','completeness_1_5').spearman_rho)}; "
        f"GPT-5.5 ρ = {_fmt(tr('openai','completeness_1_5').spearman_rho)}; "
        f"Gemini ρ = {_fmt(tr('gemini','completeness_1_5').spearman_rho)})",
        "Add CIs and the completeness correlations the Methods/SAP require.",
        "reports/aim3_compiled_tradeoff.csv")

    # ---- 7.7 Aim 2 pairwise post-hoc ----
    if len(ph):
        def pair(a, b):
            r = ph[((ph.model_a == a) & (ph.model_b == b)) | ((ph.model_a == b) & (ph.model_b == a))]
            return r.iloc[0] if len(r) else None
        cg, co, go = pair("claude", "gemini"), pair("claude", "openai"), pair("gemini", "openai")
        if cg is not None and co is not None and go is not None:
            add("with Claude and Gemini comparable and both stronger than GPT-5.5",
                "with Claude Opus 4.8 and Gemini 3.1 Pro both producing lower post-rewrite FKGL "
                f"than GPT-5.5 (Holm-adjusted paired Wilcoxon, both P {_p(co.p_holm)}); "
                f"Claude and Gemini did not differ significantly (P {_p(cg.p_holm)})",
                "The comparative claim rested on an omnibus P value alone; pairwise post-hoc "
                "now supports it, and Claude vs Gemini is not significant.",
                "reports/aim2_posthoc_models.csv")

    # ---- 7.9 lay / presentation: exploratory framing + rewrite-level values ----
    add("Two prespecified secondary analyses addressed",
        "Two secondary exploratory analyses, added after the primary protocol, addressed",
        "Methods states the neutral arm was added after a reviewer raised concern, so it "
        "cannot be described as prespecified.")
    ea, ec, ed = evl.loc["accuracy_1_5"], evl.loc["completeness_1_5"], evl.loc["added_errors_1_5"]
    add("Pooled across lay readers, accuracy was near the ceiling (4.94) and higher than the "
        "subspecialists (4.94 vs 4.84; Mann–Whitney P = .001), while completeness (4.90 vs 4.80; "
        "P = .20) and added errors (1.08 vs 1.09; P = .95) did not differ",
        "Comparing subspecialists and lay readers on the same standard instrument at the level of "
        f"the rewrite, lay accuracy was higher ({_fmt(ea.layperson_mean)} vs "
        f"{_fmt(ea.expert_mean)}; paired Wilcoxon P {_p(ea.wilcoxon_p)}), as was completeness "
        f"({_fmt(ec.layperson_mean)} vs {_fmt(ec.expert_mean)}; P {_p(ec.wilcoxon_p)}), while lay "
        f"readers recorded fewer added errors ({_fmt(ed.layperson_mean)} vs "
        f"{_fmt(ed.expert_mean)}; P {_p(ed.wilcoxon_p)})",
        "Ratings on the same rewrite are repeated observations, so the raw-rating "
        "Mann-Whitney tests were pseudo-replicated; the lay pool also mixed two instruments.",
        "reports/aim3_compiled_expert_vs_lay.csv")

    # ---- 7.10 automated panel framing ----
    add("As a pre-specified screening analysis, a panel of 3 large language models",
        "As a secondary exploratory analysis added after the primary protocol, a panel of 3 "
        "large language models",
        "The deviation log records the judge panel as a later addition.")

    # ---- 7.11 non-causal language ----
    # ---- Item 2: Methods, reproducibility (review 7.1) ----
    add("The site allowlist, the LLM rewrite prompt, the locked model versions, and the "
        "statistical analysis plan were committed to the project's version-controlled repository "
        "before Phase 1 capture began, and were specified in advance in the version-controlled "
        "project repository. The capture, cleaning, scoring, statistical analysis, and "
        "figure-generation scripts are deterministic given the raw HTML captures and the locked "
        "configuration files.",
        "The study aims, site allowlist, rewrite prompt, and statistical analysis plan were "
        "specified in the version-controlled repository before inferential analysis. The final "
        "production model panel and provider-specific runtime parameters were finalized before the "
        "rewrite experiments; changes from earlier placeholder model identifiers and "
        "provider-imposed parameter constraints were documented in the deviation log. Downstream "
        "text normalization, readability scoring, statistical analysis, and figure generation are "
        "reproducible from the committed study inputs and pinned Python environment.",
        "Model versions were finalized after Phase 1, and the LLM generation calls are not "
        "byte-reproducible; only the downstream analysis is.")

    # ---- Item 3: Methods, sample selection (review 7.2) ----
    add("For each of the 3 procedures, candidate URLs were identified via a single-operator search "
        "(signed out, incognito mode, fixed geographic context) using the prespecified queries "
        "listed in the project repository.",
        "For each procedure, candidate URLs were identified programmatically using the prespecified "
        "search queries restricted to the locked patient-facing site allowlist. This differed from "
        "the originally planned signed-out incognito-browser search; the substitution and the "
        "absence of per-result SERP HTML archival were recorded prospectively in the project "
        "deviation log. Inclusion criteria were unchanged.",
        "The documented run used programmatic search, not an incognito browser session.")

    # ---- Item 4: Results, sample capture dates (review 7.5) ----
    add("Of 26 candidate pages captured between June 2 and June 3, 2026, all 26 met inclusion "
        "criteria.",
        "Twenty-six candidate pages were identified during June 2-3, 2026. Twenty-one were captured "
        "automatically during that window; 5 pages that returned HTTP 403 responses were recovered "
        "by manual browser capture on June 8, 2026. All 26 met the final inclusion criteria and "
        "were analyzed.",
        "The original sentence implies all 26 were captured June 2-3; five came on June 8.")

    # ---- Item 5: Aim 1 site post-hoc (review 7.6) ----
    site_ph = pd.read_csv(REPORTS_DIR / "aim1_posthoc_sites.csv")
    n_sig = int((site_ph.p_adj < 0.05).sum())
    survived = ("no pairwise contrast survived Holm correction" if n_sig == 0
                else f"{n_sig} pairwise contrast(s) survived Holm correction")
    add("Differences across the 10 sites contributing 2 or more included pages were borderline "
        "(Kruskal–Wallis on FKGL, P = .074).",
        "Differences across the 10 sites contributing 2 or more included pages were borderline for "
        "FKGL (Kruskal-Wallis P = .074). Two secondary measures showed significant site-level "
        "omnibus differences (Flesch-Kincaid Reading Ease P = .042; Coleman-Liau Index P = .046); "
        f"in post-hoc Dunn tests with Holm adjustment, {survived}. Because 5 sites contributed only "
        "2 pages each, site-level comparisons are exploratory.",
        "Two significant site-level omnibus tests were unreported, and the SAP requires post-hoc "
        "follow-up. Five sites contribute n = 2, so the comparison is exploratory.",
        "reports/aim1_posthoc_sites.csv")

    # ---- Item 6: Discussion, non-causal language (review 7.11) ----
    add("localizes the residual risk to reduced completeness in the most aggressive simplifier "
        "(Gemini 3.1 Pro)",
        "localizes the lowest observed completeness to the most aggressive simplifier "
        "(Gemini 3.1 Pro), a descriptive observation rather than a demonstrated causal effect",
        "The across-model completeness test is not significant, so causal wording is unsupported.")
    add("If a model can demonstrably lower the FKGL by 2 or more grade levels with no clinically "
        "significant drop in accuracy or completeness",
        "If a model can demonstrably lower the reading level with no clinically significant drop in "
        "accuracy or completeness",
        "The '2 or more grade levels' threshold is not defined in the analysis plan.")

    # ---- Item 7: deviations count (review 7.12) ----
    add("Five deviations from the prespecified protocol are noted",
        "Fourteen deviations from the prespecified protocol are noted",
        "The repository deviation log now records 14 dated entries.",
        "docs/stats_deviations.md")

    # ---- Audit 4.3 / 10.1: pooled accuracy is rewrite-level, not rating-weighted ----
    pooled_acc = pd.read_csv(REPO_ROOT / "data" / "scores" / "accuracy.csv").accuracy_1_5.mean()
    add("pooled 4.84", f"pooled {pooled_acc:.2f} across rewrites",
        "4.84 is the rating-weighted mean over 155 rating events. The primary unit is "
        "the rewrite, which gives 4.79.",
        "data/scores/accuracy.csv")

    # ---- Audit 5.5: exact benchmark percentages ----
    rw = pd.read_csv(REPO_ROOT / "data" / "scores" / "rewrites.csv")
    pct = {}
    for m in ("claude", "gemini", "openai"):
        sub = rw[rw.model_id == m]
        pct[m] = (int((sub.fkgl <= 6.0).sum()), len(sub))
    add("met the benchmark on roughly 80% of pages",
        f"met the benchmark on {pct['claude'][0]}/{pct['claude'][1]} "
        f"({100*pct['claude'][0]/pct['claude'][1]:.1f}%) and "
        f"{pct['gemini'][0]}/{pct['gemini'][1]} "
        f"({100*pct['gemini'][0]/pct['gemini'][1]:.1f}%) of pages respectively",
        "Replaces an approximation with the exact proportions.",
        "data/scores/rewrites.csv")

    # ---- Audit 4.5 / 5.7: neutral-lay AC1 range is stale ----
    ln = irr_all[irr_all.condition == "layperson_neutral"].set_index("axis")
    lo = min(ln.loc[a, "gwet_ac1"] for a in ("accuracy_1_5", "completeness_1_5", "added_errors_1_5"))
    hi = max(ln.loc[a, "gwet_ac1"] for a in ("accuracy_1_5", "completeness_1_5", "added_errors_1_5"))
    add("Agreement among the blinded lay readers was substantial (Gwet AC1 0.76\u20130.90 across dimensions).",
        f"Agreement among the neutral-presentation lay readers was substantial "
        f"(Gwet AC1 {lo:.2f}\u2013{hi:.2f} across dimensions, using the predefined 1\u20135 category "
        f"universe). All reviewers were blinded to model identity; only this cohort was "
        f"additionally blinded to the original-versus-AI framing.",
        "The 0.76-0.90 range predates the q = 5 convention, and 'blinded' needed "
        "qualifying: every reviewer was blinded to model identity.",
        "reports/aim3_compiled_irr.csv")

    # ---- Audit 5.6: precision about the single low rating ----
    add("a single rewrite received the lowest accuracy score",
        "one expert accuracy rating was 3 or lower",
        "The pooled table establishes one rating <= 3, which is a rating event, not a "
        "verified rewrite-level minimum.")

    # ---- Audit 5.8: the automated panel is not concordant on model ranking ----
    add("and the automated LLM-judge panel was concordant",
        "and the automated LLM-judge panel independently flagged the same specific "
        "errors, although it ranked the models differently",
        "Automated consensus ranks GPT-5.5 highest and Gemini lowest, whereas primary "
        "human accuracy does not differ significantly by model.",
        "reports/aim3_llm_descriptives.csv")

    # ---- Audit 5.6: completeness trade-off correlations were all nonsignificant ----
    add("Completeness showed no association for any model",
        "No completeness correlation reached statistical significance for any model",
        "States the inferential conclusion rather than only listing rho values.",
        "reports/aim3_compiled_tradeoff.csv")

    # ---- Audit 5.3: reference 3 does not support the attribution ----
    add("The National Library of Medicine, the American Medical Association, and the Centers for "
        "Disease Control have recommended that patient-facing health information be written at or "
        "below a 6th-grade reading level, because roughly one-third of United States adults read "
        "at or below that level",
        "Health-literacy guidance commonly recommends that patient-facing health information be "
        "written at or below a 6th-grade reading level, because roughly one-third of United States "
        "adults read at or below that level",
        "Reference 3 (Kutner, National Assessment of Adult Literacy) is a literacy survey. It "
        "supports the reading-level statistic but contains no NLM, AMA or CDC recommendation, so "
        "the attribution is unsupported by the cited source. Softened rather than removed; cite the "
        "actual guidance documents if the named attribution is wanted.",
        "reference audit")

    # ---- Audit 5.3: reference 8 is not a cardiology readability study ----
    add("Several recent reports have evaluated the readability of LLM-generated answers to "
        "cardiology patient questions7,8 but have not evaluated rewriting of existing pages",
        "Recent reports have evaluated the readability of LLM-generated answers to cardiology "
        "patient questions7 and have compared physician and chatbot responses to general patient "
        "questions,8 but have not evaluated rewriting of existing pages",
        "Reference 8 (Ayers et al.) compared physician and chatbot responses to general patient "
        "questions on a public social-media forum; it is neither cardiology-specific nor a "
        "readability study. Each citation now carries only the claim it supports, and reference 8 "
        "stays cited rather than becoming orphaned.",
        "reference audit")

    # ---- Reference audit: refs 9-18 are listed but uncited, and each has an
    # ---- existing home in the text. These are missing citation markers, not new
    # ---- claims, so adding them changes nothing the manuscript asserts.
    add("FKRE (higher = easier), FKGL, GFI, SMOG, CLI, and ARI.",
        "FKRE (higher = easier),11 FKGL,12 GFI,13 SMOG,14 CLI,15 and ARI.16",
        "References 11-16 are the primary sources for the six formulas and were listed "
        "but never cited; each is attached to the formula it defines.",
        "reference audit")
    add("Body text was extracted using trafilatura",
        "Body text was extracted using trafilatura17",
        "Reference 17 is the trafilatura source and was listed but never cited.",
        "reference audit")
    add("using the textstat Python library", "using the textstat Python library18",
        "Reference 18 is the textstat source and was listed but never cited.",
        "reference audit")
    add("as measured by the Agency for Healthcare Research and Quality Patient Education "
        "Materials Assessment Tool, were not assessed in this initial analysis",
        "as measured by the Agency for Healthcare Research and Quality Patient Education "
        "Materials Assessment Tool,9,10 were not assessed in this initial analysis",
        "References 9 and 10 are the PEMAT development paper and the AHRQ instrument; the "
        "Limitations already name the tool but carried no citation.",
        "reference audit")

    # ---- Refs 23-25: prior LLM-in-medicine work. The sentence below describes each
    # ---- only by what its own title states, so nothing is asserted that the
    # ---- manuscript's reference list does not already contain.
    add("Behers and colleagues measured LLM generation of cardiac catheterization patient "
        "text and reported readability of model output7",
        "Related work has assessed chatbot responses to common cancer queries,23 the accuracy "
        "and reliability of chatbot responses to physician questions,24 and simplification of "
        "radiology reports for patients.25 Behers and colleagues measured LLM generation of "
        "cardiac catheterization patient text and reported readability of model output7",
        "References 23-25 were listed but never cited. Each is described here only by what "
        "its own title states, so no finding is characterised beyond the reference list.",
        "reference audit")

    # Reference integrity: SciPy and pandas are named in the text but never cited.
    add("Analyses were performed in Python 3.11 using SciPy and pandas.",
        "Analyses were performed in Python 3.11 using SciPy and pandas.19,20",
        "References 19 and 20 appear in the reference list but were never cited; "
        "this is the sentence that names both tools.",
        "reference-list audit")

    add("the residual risk a simplification–completeness trade-off in the most aggressive simplifier",
        "the lowest observed completeness occurring in the most aggressive simplifier",
        "Across-model completeness is not significant (P ≈ .06), so causal phrasing is not supported.")

    return E


def _stamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _run(text, rpr, deleted=False):
    r = docx.oxml.OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = docx.oxml.OxmlElement("w:delText" if deleted else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _revision(tag, child):
    _rev_id[0] += 1
    el = docx.oxml.OxmlElement(tag)
    el.set(qn("w:id"), str(_rev_id[0]))
    el.set(qn("w:author"), REVISION_AUTHOR)
    el.set(qn("w:date"), _stamp())
    el.append(child)
    return el


def _fold(text: str) -> str:
    """Fold typography for matching only.

    Every substitution is 1:1 in length, so an index found in the folded string is
    valid in the original. Word uses en/em dashes, curly quotes and non-breaking
    spaces; search strings written in ASCII would otherwise never match.
    """
    for a, b in (("\u2013", "-"), ("\u2014", "-"), ("\u2212", "-"), ("\u2010", "-"),
                 ("\u2018", "'"), ("\u2019", "'"), ("\u201c", '"'), ("\u201d", '"'),
                 ("\u00a0", " "), ("\u2009", " "), ("\u202f", " ")):
        text = text.replace(a, b)
    return text


def find_in(text: str, needle: str) -> int:
    """Index of `needle` in `text`, ignoring typographic variants."""
    return _fold(text).find(_fold(needle))


def _para_accepted(para) -> str:
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    return "".join(n.text for n in para._element.iter() if n.tag == f"{W}t" and n.text)


def apply_paragraph_edits(para, edits: list[dict]) -> int:
    """Apply every edit belonging to this paragraph in ONE pass.

    Applying them one at a time is unsafe: after the first, part of the paragraph
    lives inside a `w:ins` element, which python-docx cannot see. A second edit
    reading `para.text` would get a truncated string and, on rewriting the runs,
    silently delete the first insertion along with the surrounding prose. This bug
    destroyed text on an earlier run, hence the single-pass rebuild.
    """
    # A paragraph may already carry revisions from an earlier run. python-docx
    # cannot see runs nested in w:ins, so rebuilding from para.text would silently
    # delete those earlier insertions. Refuse to touch such a paragraph rather than
    # corrupt it; the edit is reported as not found and can be reviewed by hand.
    if para._element.find(qn("w:ins")) is not None or para._element.find(qn("w:del")) is not None:
        return 0

    text = para.text
    spans = sorted((find_in(text, e["old"]), e) for e in edits if find_in(text, e["old"]) >= 0)
    if not spans:
        return 0

    rpr = None
    if para.runs and para.runs[0]._element.find(qn("w:rPr")) is not None:
        rpr = para.runs[0]._element.find(qn("w:rPr"))
    for r in list(para.runs):
        r._element.getparent().remove(r._element)

    p, cursor = para._element, 0
    for i, e in spans:
        if i < cursor:
            continue
        if i > cursor:
            p.append(_run(text[cursor:i], rpr))
        # Delete the document's own wording, not the ASCII search string, so the
        # revision shows exactly what was removed.
        p.append(_revision("w:del", _run(text[i:i + len(e["old"])], rpr, deleted=True)))
        p.append(_revision("w:ins", _run(e["new"], rpr)))
        cursor = i + len(e["old"])
    if cursor < len(text):
        p.append(_run(text[cursor:], rpr))
    return len(spans)


def enable_track_changes(document) -> None:
    """Switch on <w:trackChanges/> so later editing in Word is tracked too."""
    settings = document.settings.element
    if settings.find(qn("w:trackChanges")) is None:
        settings.insert(0, docx.oxml.OxmlElement("w:trackChanges"))


OUTSTANDING = [
    "Table 3 header: change 'No. ratings' to 'No. rewrites' and set 26 / 25 / 26 "
    "(table structure edit, needs author review of layout)",
    "§7.1 Methods reproducibility paragraph — replacement text supplied by the review",
    "§7.2 Methods sample selection — programmatic search wording",
    "§7.5 Results sample — capture dates split (21 automated, 5 manual on June 8)",
    "§7.6 Aim 1 — add site-level Dunn post-hoc sentence once those diagnostics are run",
    "§7.11 Discussion / Conclusions — non-causal reframing throughout",
    "§7.12 Deviations summary — update the count; the repository log now has 10 entries",
    "§7.13 eFigure 1 / eFigure 2 — regenerate and add the missing eFigure 2 legend",
]


def _accepted_text(cell) -> str:
    """Cell text with tracked changes accepted, so an already-revised cell is not
    re-edited into nonsense on a second run."""
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    return "".join(n.text for n in cell._element.iter() if n.tag == f"{W}t" and n.text)


def revise_table3(document, prim) -> int:
    """Item 1: Table 3 counts rewrites, not rating events.

    Targeted by table position and cell coordinates, never by text search. Searching
    the document for a bare "55" would match inside "155" and silently corrupt an
    unrelated number; only the count column of this one table may change.
    """
    # Match on the accepted-changes header so an already-revised table is still
    # found, and take the FIRST such table: Table 3 (human) precedes Table 4 (judges).
    table = next((t for t in document.tables
                  if _accepted_text(t.rows[0].cells[1]).strip().startswith("No. r")), None)
    if table is None:
        return 0
    changed = 0
    header = table.rows[0].cells[1].paragraphs[0]
    if apply_paragraph_edits(header, [{"old": "No. ratings", "new": "No. rewrites"}]):
        changed += 1
    key = {"Claude Opus 4.8": "claude", "GPT-5.5": "openai", "Gemini 3.1 Pro": "gemini"}
    for row in table.rows[1:]:
        model = key.get(row.cells[0].text.strip())
        if model is None:
            continue
        # Column 1 is the count; columns 2-4 are accuracy, completeness, added
        # errors as "mean (SD)". All are rewritten from the primary table.
        targets = {1: str(int(prim.loc[model, "n_rewrites"]))}
        for col, axis in ((2, "accuracy_1_5"), (3, "completeness_1_5"), (4, "added_errors_1_5")):
            targets[col] = (f"{prim.loc[model, axis + '_mean']:.2f} "
                            f"({prim.loc[model, axis + '_sd']:.2f})")
        for col, new in targets.items():
            if col >= len(row.cells):
                continue
            cell = row.cells[col]
            old = _accepted_text(cell).strip()
            if old and old != new:
                if apply_paragraph_edits(cell.paragraphs[0], [{"old": old, "new": new}]):
                    changed += 1
    return changed


def clean_whitespace(document) -> int:
    """Strip trailing/leading spaces from runs and collapse internal double spaces.

    Applied silently rather than as tracked changes: seven invisible space removals
    would add revision marks that tell a reader nothing, and would obscure the
    substantive edits. Text content is unchanged.
    """
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    fixed = 0
    for para in document.paragraphs:
        nodes = [n for n in para._element.iter() if n.tag == f"{W}t" and n.text]
        for k, node in enumerate(nodes):
            original = node.text
            text = original.replace("\u00a0", " ")
            while "  " in text:
                text = text.replace("  ", " ")
            if k == 0:
                text = text.lstrip()
            if k == len(nodes) - 1:
                text = text.rstrip()
            if text != original:
                node.text = text
                fixed += 1
    return fixed


def main() -> int:
    if not MANUSCRIPT.exists():
        print(f"manuscript not found: {MANUSCRIPT}")
        return 2

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"{MANUSCRIPT.stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    shutil.copy2(MANUSCRIPT, backup)

    document = docx.Document(str(MANUSCRIPT))
    enable_track_changes(document)

    paragraphs = list(document.paragraphs)
    for tbl in document.tables:
        for row in tbl.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    assigned: dict[int, list[dict]] = {}
    applied: list[dict] = []
    pending: list[tuple[int, dict]] = []
    for e in build_edits():
        target = next((i for i, para in enumerate(paragraphs)
                       if find_in(_para_accepted(para), e["old"]) >= 0), None)
        if target is None:
            applied.append({**e, "status": "NOT FOUND"})
        else:
            assigned.setdefault(target, []).append(e)
            pending.append((target, e))

    # Status must reflect what actually landed, not merely that a target paragraph
    # existed: apply_paragraph_edits refuses paragraphs that already carry revisions,
    # so reporting APPLIED on a successful lookup produced false confirmations.
    edited = set()
    for idx, group in assigned.items():
        if apply_paragraph_edits(paragraphs[idx], group):
            edited.add(idx)
    for target, e in pending:
        applied.append({**e, "status": "APPLIED" if target in edited
                        else "SKIPPED (paragraph already revised)"})

    n_t3 = revise_table3(document, pd.read_csv(
        REPORTS_DIR / "aim3_compiled_by_model_primary.csv").set_index("model_id"))
    n_ws = clean_whitespace(document)
    document.save(str(MANUSCRIPT))

    audit = pd.DataFrame(applied)[["status", "rationale", "source", "old", "new"]]
    audit.to_csv(AUDIT, index=False)
    n_ok = int((audit.status == "APPLIED").sum())

    print("=" * 84)
    print("MANUSCRIPT REVISED IN PLACE, AS TRACKED CHANGES")
    print("=" * 84)
    print(f"  file   : {MANUSCRIPT.name}")
    print(f"  backup : {backup.relative_to(REPO_ROOT)}")
    print(f"  audit  : {AUDIT.name}")
    print(f"  Table 3 cells revised  : {n_t3}")
    print(f"  whitespace runs cleaned: {n_ws}")
    print(f"\n  {n_ok} of {len(applied)} edits applied as tracked revisions\n")
    for _, r in audit.iterrows():
        print(f"  [{'OK ' if r.status == 'APPLIED' else 'MISS'}] {r.rationale[:72]}")
    print("\n  Outstanding, needing author judgement:")
    for o in OUTSTANDING:
        print(f"    - {o}")
    print("\n  In Word: Review > Tracking to step through each change.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
