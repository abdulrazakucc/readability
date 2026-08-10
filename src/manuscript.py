"""Extract every reported value from Dr Naeem's manuscript .docx.

Why parse rather than transcribe
--------------------------------
The validation harness needs an *independent* reference to check the pipeline
against. Two tempting shortcuts are both wrong:

* Reading expected values from our own `reports/` would be circular -- the check
  would pass by construction and certify nothing.
* Transcribing the numbers into Python by hand makes them stale the moment the
  manuscript is revised, and a transcription slip looks exactly like a pipeline
  bug.

So we read them straight out of the .docx. The manuscript stays the source of
truth, revisions are picked up automatically, and nothing is duplicated.

What this module does NOT do
----------------------------
It does not compute anything and it does not know what any number *should* be.
It only reports what the document says. The comparison lives in
`scripts/14_validate_manuscript.py`.

Every extracted value carries the raw string it came from, so a mismatch can be
traced back to the exact sentence or table cell.
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path

import docx

# Formula names as printed in Table 1 / Table 2, mapped to our score columns.
# This is a label mapping, not a value: the numbers still come from the document.
FORMULA_KEYS: dict[str, str] = {
    "flesch-kincaid reading ease": "fkre",
    "flesch-kincaid grade level": "fkgl",
    "gunning fog index": "gfi",
    "smog index": "smog",
    "coleman-liau index": "cli",
    "automated readability index": "ari",
}

MODEL_KEYS: dict[str, str] = {
    "claude opus 4.8": "claude",
    "gpt-5.5": "openai",
    "gemini 3.1 pro": "gemini",
}


@dataclass(frozen=True)
class Reported:
    """A single value as printed in the manuscript.

    `decimals` drives the comparison tolerance: a value printed as "4.84" is only
    claimed to 2 decimal places, so it is checked to +/-0.005. That way the
    tolerance follows the document's own precision instead of being chosen by us.
    """

    key: str
    value: float
    raw: str
    decimals: int
    source: str  # "Table 2" / "Results" etc, for traceability

    @property
    def tolerance(self) -> float:
        return 0.5 * (10 ** -self.decimals)


def normalize(text: str) -> str:
    """Fold the typographic characters a Word document is full of.

    Word uses U+2212 MINUS SIGN, en/em dashes, non-breaking spaces, Greek letters
    and superscript digits. Regexes written against ASCII silently fail on all of
    them, so everything is normalised once, here.
    """
    t = unicodedata.normalize("NFKC", text)
    for bad, good in [
        ("−", "-"), ("–", "-"), ("—", "-"), ("‐", "-"), ("‑", "-"),
        (" ", " "), (" ", " "), (" ", " "),
        ("‘", "'"), ("’", "'"), ("“", '"'), ("”", '"'),
        ("×", "x"), ("ρ", "rho"), ("κ", "kappa"), ("χ", "chi"),
        ("≤", "<="), ("≥", ">="),
    ]:
        t = t.replace(bad, good)
    # Superscript digits (10^-9 written as 10⁻⁹) -> caret form.
    sup = {"⁰": "0", "¹": "1", "²": "2", "³": "3", "⁴": "4",
           "⁵": "5", "⁶": "6", "⁷": "7", "⁸": "8", "⁹": "9",
           "⁻": "-"}
    out = []
    prev_sup = False
    for ch in t:
        if ch in sup:
            if not prev_sup:
                out.append("^")
            out.append(sup[ch])
            prev_sup = True
        else:
            out.append(ch)
            prev_sup = False
    return re.sub(r"\s+", " ", "".join(out))


def _decimals(s: str) -> int:
    m = re.search(r"\.(\d+)", s.strip())
    return len(m.group(1)) if m else 0


def _num(s: str) -> float:
    """Parse a printed number, including the bare-point form JAMA uses for P (.06)."""
    s = s.strip().replace(",", "")
    if s.startswith("."):
        s = "0" + s
    elif s.startswith("-."):
        s = "-0" + s[1:]
    return float(s)


def _key(text: str) -> str:
    """Lowercase, dash-folded lookup key for a table row label."""
    t = normalize(text).lower().strip()
    t = re.sub(r"\s*\(.*?\)\s*", "", t)          # drop "(higher = easier)"
    return re.sub(r"[^a-z0-9. -]", "", t).strip()


def _add(out: dict[str, Reported], key: str, raw: str, source: str) -> None:
    out[key] = Reported(key=key, value=_num(raw), raw=raw.strip(),
                        decimals=_decimals(raw), source=source)


# --------------------------------------------------------------------------
# Tables
# --------------------------------------------------------------------------

def parse_tables(doc: docx.document.Document) -> dict[str, Reported]:
    """Read Tables 1-4. Tables are identified by their header row, not position,
    so inserting a table upstream does not silently shift the mapping."""
    out: dict[str, Reported] = {}
    for t in doc.tables:
        rows = [[normalize(c.text) for c in r.cells] for r in t.rows]
        if not rows:
            continue
        header = " | ".join(rows[0]).lower()

        # Table 1: formula | No. | Mean | SD | Median
        if "readability formula" in header and "mean" in header and "median" in header:
            for r in rows[1:]:
                f = FORMULA_KEYS.get(_key(r[0]))
                if not f:
                    continue
                _add(out, f"t1.{f}.n", r[1], "Table 1")
                _add(out, f"t1.{f}.mean", r[2], "Table 1")
                _add(out, f"t1.{f}.sd", r[3], "Table 1")
                _add(out, f"t1.{f}.median", r[4], "Table 1")

        # Table 2: formula | <model> delta (95% CI) x3 | Friedman chi2 (P)
        elif "readability formula" in header and "friedman" in header:
            cols = {}
            for i, h in enumerate(rows[0][1:], start=1):
                for label, mk in MODEL_KEYS.items():
                    if label in h.lower():
                        cols[i] = mk
            for r in rows[1:]:
                f = FORMULA_KEYS.get(_key(r[0]))
                if not f:
                    continue
                for i, mk in cols.items():
                    m = re.search(r"(-?[\d.]+)\s*\(\s*(-?[\d.]+)\s*to\s*(-?[\d.]+)\s*\)", r[i])
                    if m:
                        _add(out, f"t2.{f}.{mk}.delta", m.group(1), "Table 2")
                        _add(out, f"t2.{f}.{mk}.ci_low", m.group(2), "Table 2")
                        _add(out, f"t2.{f}.{mk}.ci_high", m.group(3), "Table 2")
                chi = re.search(r"([\d.]+)\s*\(", r[-1])
                if chi:
                    _add(out, f"t2.{f}.friedman_chi2", chi.group(1), "Table 2")

        # Tables 3 and 4 share a shape; the count column tells them apart.
        elif "accuracy" in header and "completeness" in header and "added errors" in header:
            tag = "t4" if "rewrites" in header else "t3"
            src = "Table 4" if tag == "t4" else "Table 3"
            for r in rows[1:]:
                mk = MODEL_KEYS.get(_key(r[0]))
                if not mk:
                    continue
                _add(out, f"{tag}.{mk}.n", r[1], src)
                for j, axis in enumerate(["accuracy", "completeness", "added_errors"], start=2):
                    m = re.search(r"([\d.]+)\s*\(([\d.]+)\)", r[j])
                    if m:
                        _add(out, f"{tag}.{mk}.{axis}.mean", m.group(1), src)
                        _add(out, f"{tag}.{mk}.{axis}.sd", m.group(2), src)
    return out


# --------------------------------------------------------------------------
# Prose
# --------------------------------------------------------------------------

# Each entry: key template, regex, and the group(s) to capture. Patterns are
# anchored on distinctive wording so they fail loudly (missing key) rather than
# matching the wrong sentence.
_N = r"(-?[\d.]+)"

_PROSE: list[tuple[str | list[str], str]] = [
    # Aim 1
    (["p.fkgl.median", "p.fkgl.mean", "p.fkgl.sd", "p.fkgl.min", "p.fkgl.max"],
     rf"median FKGL was {_N} \(mean {_N}; SD {_N}; range, {_N}-{_N}\)"),
    (["p.benchmark.meeting", "p.benchmark.n", "p.benchmark.pct"],
     rf"{_N} of {_N} included pages \({_N}%"),
    (["p.benchmark.ci_high"], rf"95% CI, 0%-{_N}%"),
    (["p.words.mean", "p.words.min", "p.words.max"],
     rf"Mean cleaned body length was {_N} words \(range, {_N}-{_N}\)"),
    (["p.proc.tavr.mean", "p.proc.tavr.sd", "p.proc.tavr.n"],
     rf"TAVR, {_N} \(SD {_N}; n = {_N}\)"),
    (["p.proc.cta.mean", "p.proc.cta.sd", "p.proc.cta.n"],
     rf"CCTA, {_N} \(SD {_N}; n = {_N}\)"),
    (["p.proc.laao.mean", "p.proc.laao.sd", "p.proc.laao.n"],
     rf"LAAO, {_N} \(SD {_N}; n = {_N}\)"),
    (["p.proc.anova_p"], rf"analysis of variance F-test, P = {_N}"),
    (["p.site.kruskal_p"], rf"Kruskal-Wallis on FKGL, P = {_N}"),
    # Aim 2
    (["p.dz.gemini"], rf"Gemini 3.1 Pro reduced the grade level by a mean of [\d.]+ \(95% CI, [\d.]+-[\d.]+; Cohen d_?z = {_N}\)"),
    (["p.dz.claude"], rf"Claude Opus 4.8 by [\d.]+ \(95% CI, [\d.]+-[\d.]+; d_?z = {_N}\)"),
    (["p.dz.openai"], rf"GPT-5.5 by [\d.]+ \(95% CI, [\d.]+-[\d.]+; d_?z = {_N}\)"),
    (["p.post.gemini", "p.post.claude", "p.post.openai"],
     rf"Post-rewrite mean FKGL was {_N} for Gemini, {_N} for Claude, and {_N} for GPT-5.5"),
    (["p.met.gemini", "p.met.gemini_n", "p.met.claude", "p.met.claude_n",
      "p.met.openai", "p.met.openai_n"],
     rf"to {_N} of {_N} \(Gemini\), {_N} of {_N} \(Claude\), and {_N} of {_N} \(GPT-5.5\)"),
    (["p.friedman.fkgl_chi2", "p.friedman.n_pages"],
     rf"Friedman chi\^?2? = {_N}; P = [\d.]+ x 10\^?-9; {_N} pages"),
    # Aim 3 primary
    (["p.aim3.n_reviewers"], rf"{_N} cardiothoracic radiologists"),
    (["p.aim3.n_ratings"], rf"contributing {_N} independent ratings"),
    (["p.aim3.pooled_accuracy"], rf"pooled {_N}\)"),
    (["p.aim3.pct_ge4", "p.aim3.pct_eq5", "p.aim3.pct_added_le2"],
     rf"{_N}% of accuracy ratings were 4 or 5 \({_N}% were 5\) and {_N}% of added-error ratings"),
    (["p.irr.pairs", "p.irr.accuracy_exact", "p.irr.completeness_exact", "p.irr.added_errors_exact"],
     rf"{_N} rater-pairs; exact agreement {_N}% for accuracy, {_N}% for completeness, and {_N}% for added errors"),
    (["p.ac1.accuracy", "p.ac1.completeness", "p.ac1.added_errors"],
     rf"Gwet AC1 {_N} for accuracy, {_N} for completeness, and {_N} for added errors"),
    (["p.kappa.accuracy", "p.kappa.completeness", "p.kappa.added_errors"],
     rf"Cohen kappa was near (?:zero|0) \({_N}, {_N}, and {_N}\)"),
    (["p.rho.gemini", "p.rho.gemini_p"],
     rf"Spearman rho = {_N} between grade levels removed and accuracy; P = {_N}\)"),
    (["p.rho.claude", "p.rho.claude_p"], rf"Claude Opus 4.8 \(rho = {_N}; P = {_N}\)"),
    (["p.rho.openai", "p.rho.openai_p"], rf"GPT-5.5 \(rho = {_N}; P = {_N}\)"),
    # Lay arm and presentation
    (["p.lay.n_neutral", "p.lay.n_labeled", "p.lay.n_ratings"],
     rf"{_N} scored the neutral-presentation variant .*? and {_N} scored the standard, labeled instrument, for {_N} lay ratings"),
    (["p.lay.accuracy", "p.lay.expert_accuracy", "p.lay.accuracy_p"],
     rf"higher than the subspecialists \({_N} vs {_N}; Mann-Whitney P = {_N}\)"),
    (["p.lay.completeness", "p.lay.expert_completeness", "p.lay.completeness_p"],
     rf"while completeness \({_N} vs {_N}; P = {_N}\)"),
    (["p.lay.added", "p.lay.expert_added", "p.lay.added_p"],
     rf"added errors \({_N} vs {_N}; P = {_N}\) did not differ"),
    # Abstract / Key Points restatements (rounded to 1dp) and site-level detail
    (["p.abs.fkgl_median", "p.abs.iqr_low", "p.abs.iqr_high"],
     rf"median (?:original )?(?:Flesch-Kincaid Grade Level|FKGL) was {_N} \(IQR, {_N}-{_N}\)"),
    (["p.abs.fkgl_median2", "p.abs.iqr_low2", "p.abs.iqr_high2"],
     rf"median (?:Flesch-Kincaid Grade Level|FKGL) was {_N} \(interquartile range, {_N}-{_N}\)"),
    (["p.fkre.mean_prose", "p.fkre.sd_prose"], rf"mean FKRE was {_N} \(SD {_N}\)"),
    (["p.gfi.mean_prose", "p.smog.mean_prose"], rf"mean GFI \({_N}\) and SMOG \({_N}\)"),
    (["p.site.mayo_mean", "p.site.mayo_n"], rf"Mayo Clinic \(mean FKGL {_N}; n = {_N}\)"),
    (["p.site.bwh_mean", "p.site.bwh_n"],
     rf"Brigham and Women's Hospital \(mean FKGL {_N}; n = {_N}\)"),
    (["p.site.bhf_mean", "p.site.radinfo_mean"],
     rf"mean FKGL {_N} \(British Heart Foundation\) to {_N} \(RadiologyInfo"),
    (["p.sample.cta_n", "p.sample.tavr_n", "p.sample.laao_n"],
     rf"CCTA, {_N} pages; TAVR, {_N} pages; LAAO, {_N} pages"),
    # Aim 3 secondary: automated LLM-judge panel
    (["p.llm.n_judgments"], rf"blinded to the producing model \({_N} judgments"),
    (["p.llm.pct_max"], rf"{_N}% of accuracy ratings were maximal and no judgment"),
    (["p.llm.gemini_accuracy"], rf"had the lowest consensus accuracy \({_N}/5\)"),
    (["p.llm.gemini_added"], rf"most added content \(added-errors {_N}/5\)"),
    (["p.llm.openai_accuracy", "p.llm.openai_added"],
     rf"was the most faithful \({_N}; {_N}\)"),
    (["p.llm.claude_accuracy"], rf"was the best balance \({_N} accuracy"),
    (["p.llm.friedman_chi2"], rf"across-model accuracy differed significantly \(Friedman chi\^?2? = {_N}"),
    (["p.llm.self_pref_own", "p.llm.self_pref_other", "p.llm.self_pref_p"],
     rf"own rewrites higher on accuracy than other models' \({_N} vs {_N}; Mann-Whitney P = {_N}\)"),
    (["p.pres.accuracy_labeled", "p.pres.accuracy_neutral",
      "p.pres.completeness_labeled", "p.pres.completeness_neutral",
      "p.pres.added_labeled", "p.pres.added_neutral"],
     rf"accuracy {_N} vs {_N}; completeness {_N} vs {_N}; added errors {_N} vs {_N}"),
]


_WORD_NUM = {"zero": "0", "one": "1", "two": "2", "three": "3", "four": "4", "five": "5",
             "six": "6", "seven": "7", "eight": "8", "nine": "9", "ten": "10"}


def _fold_word_numbers(text: str) -> str:
    """The manuscript spells some counts as words ("Zero of 26", "Six
    cardiothoracic radiologists"). Fold them to digits so one numeric pattern
    covers both styles. Applied to prose only -- tables are already numeric."""
    def sub(m):
        return _WORD_NUM[m.group(0).lower()]
    return re.sub(r"\b(" + "|".join(_WORD_NUM) + r")\b", sub, text, flags=re.IGNORECASE)


def parse_prose(doc: docx.document.Document) -> dict[str, Reported]:
    text = _fold_word_numbers(normalize(" ".join(p.text for p in doc.paragraphs)))
    out: dict[str, Reported] = {}
    for keys, pattern in _PROSE:
        m = re.search(pattern, text, flags=re.IGNORECASE)
        if not m:
            continue
        keys = [keys] if isinstance(keys, str) else keys
        for i, k in enumerate(keys, start=1):
            if i <= len(m.groups()):
                _add(out, k, m.group(i), "Results")
    return out


def load_reported(path: str | Path) -> dict[str, Reported]:
    """All values the manuscript reports, keyed for comparison."""
    doc = docx.Document(str(path))
    values = parse_tables(doc)
    values.update(parse_prose(doc))
    if not values:
        raise SystemExit(f"extracted no values from {path} — has the format changed?")
    return values
